from odoo import models, fields, api
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
import base64
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
class facturas(models.Model):
     _name ='empresa_abogados.facturas'
     _description ='Modelos para almacenar facturas'
     _inherit='empresa_abogados.base'
     
     proyecto= fields.Many2one('empresa_abogados.proyecto', string='Proyecto',required=True,ondelete='cascade')

     fecha = fields.Date(string='Fecha Factura',default=fields.Datetime.now)    
     num_factura = fields.Char(string='Numero de factura')
     concepto = fields.Text(string='Concepto')
     
     tramitacion = fields.Selection(related='proyecto.tramitacion', string='Tramitación')
     BI = fields.Float(string='Base Imponible')
     IVA= fields.Float(string='IVA',default=0.21)
     importe_IVA= fields.Float(compute='calculo_importe_IVA',string='Importe IVA')
   
     importe_total = fields.Float(compute='calculo_importe',string='Importe Total')

     fecha_cobro = fields.Date(string='Fecha cobro',default=fields.Datetime.now)    
   
     cobrado = fields.Boolean(string='Cobrado')
     Modo_envio = fields.Selection([('0','Ordinario'),('1','Face'),('2','Equipo')],string='Modo de Envio')
    
    
     @api.depends('IVA','BI','importe_IVA')
     def calculo_importe_IVA(self):
        for elemento in self:
            bi= elemento.BI
            iva= elemento.IVA
            elemento.importe_IVA= (bi * iva)
     cliente = fields.Many2one(related='proyecto.cliente', string='Cliente',required=True)#revisar
     @api.depends('BI','IVA','importe_total')
     def calculo_importe(self):
        for elemento in self:
           
            elemento.importe_total= elemento.importe_IVA + elemento.BI

     @api.depends('cliente')
     def _compute_cliente_data(self):
        for factura in self:
            cliente = factura.cliente
            if cliente:
                factura.nombre_cliente = cliente.name
                factura.direccion_cliente = cliente.direccion
    #datos reporte
     nombre_cliente = fields.Char(string='Nombre del cliente', compute='_compute_cliente_data', store=True)
     direccion_cliente = fields.Text(string='Dirección del cliente', compute='_compute_cliente_data', store=True)
     cif_cliente = fields.Char(compute="_compute_cliente_cif_data",string="Cif cliente ")
     @api.depends('cliente')
     def _compute_cliente_cif_data(self):
        for factura in self:
            cliente = factura.cliente
            if cliente:
                factura.cif_cliente = cliente.cif
    
     def reporte_facturas_arca_consortium (self):

        doc = Document()
        section = doc.sections[0]
        header = section.header
        paragraph = header.add_paragraph('')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Agregar la imagen al párrafo
        image_path = "/var/lib/odoo/modulo/empresa_abogados/static/src/img/arca_consortium.png"  # Reemplaza con la ruta de tu imagen
        run = paragraph.add_run()
        run.add_picture(image_path,width=Inches(2.5))  # Ajusta el ancho según tus necesidades

       
        doc.add_paragraph('FACTURA '+str(self.num_factura)).alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph('')
        
        # Add the client 
        doc.add_paragraph(str(self.cliente.name))
        direccion = self.cliente.direccion if self.cliente.direccion else ""
        doc.add_paragraph(str(direccion))
        cif= self.cliente.cif if self.cliente.cif else ""
        doc.add_paragraph(str(f'C.I.F.'+ cif))
        doc.add_paragraph('')
        # Add the project name
        doc.add_paragraph('Factura correspondiente a los servicios '+str(self.proyecto.name))

        # Add the invoice details
        doc.add_paragraph('')
        doc.add_paragraph('Base Imponible .................................... ' + str(self.BI)+'€').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph('% IVA ............................................. ' + str( round(self.importe_IVA, 2)) +'€').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph('TOTAL ............................................. ' + str( round(self.importe_total, 2)) +'€').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph('')
        # Add the signature and date
        doc.add_paragraph('')
        doc.add_paragraph('En Madrid, a fecha y hora de la firma electrónica').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph('')
        doc.add_paragraph('')
        doc.add_paragraph('Fdo. Esther Camacho González').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph('')
        # Add the payment information
        doc.add_paragraph('')
        doc.add_paragraph(' Dicha factura la pueden hacer efectiva mediante transferencia a nuestra cuenta corriente abierta en La Caixa (Calle Génova 17 de Madrid), ES3721009194132200582579 a nombre de ARCA CONSORTIUM S.A.').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add the data protection information
        doc.add_paragraph('')
        doc.add_paragraph('De conformidad con lo establecido en la normativa vigente en Protección de Datos de Carácter Personal, le informamos que sus datos serán incorporados al sistema de tratamiento titularidad de ARCA CONSORTIUM SA  con CIF A28911238 y domicilio social sito en CL SANTA ENGRACIA 6 28010, MADRID, con la finalidad de poder remitirle la correspondiente factura. En cumplimiento con la normativa vigente, ARCA CONSORTIUM SA  informa que los datos serán conservados durante EN EL PLAZO LEGALMENTE ESTABLECIDO.')
        doc.add_paragraph('(Calle Arturo Soria 316 de Madrid). C.C.C. nº ES4100810638590002108117 a nombre de ARCA AUDITORES S.L.')
       
        section = doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Agregar el texto al footer con el formato deseado
        footer_text = 'SANTA ENGRACIA 6 – 28010 MADRID – TEL.:34-91 319 26 58 – FAX: 34 – 91 319 27 51  e-mail: madrid@arcaconsortium.net\n'
        footer_text += 'Rue Belliard 20 – 1040 BRUXELES – e-mail : bruselas@arcaconsortium.net\n'
        footer_text += 'ARCA CONSORTIUM, S.A. Inscrita en R.M. MM Madrid, tomo 7108, folio 58, sección 8ª, hoja M-115.454. C.I.F. A-28911238'
        # Agregar el texto al párrafo y establecer el tamaño de fuente y color
        run = paragraph.add_run(footer_text)
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(25, 25, 112)  # Color hexadecimal #191970
        run.font.bold=True
        # Save the document
        doc.save('factura.docx')
        # Convert the saved document to binary data
        with open('factura.docx', 'rb') as f:
            binary_data = f.read()

        # Create the attachment record
        attachment = self.env['ir.attachment'].sudo().create({
            'datas': base64.encodestring(binary_data),
            'type': 'binary',
            'res_model': 'empresa_abogados.facturas',
            'res_id': self.proyecto.id,
            'db_datas': 'facturas_Arca_Consortium.docx',
            'name': 'facturas_Arca_Consortium.docx',
        })
        return {
            'type': 'ir.actions.act_url',
            'url': '/web/content/%s?download=true' % attachment.id,
            'target': 'new',
        }

     def reporte_facturas_arca_auditores(self):
        
        doc = Document()
        
        section = doc.sections[0]
        header = section.header
        paragraph = header.add_paragraph('')
        # Agregar la imagen al párrafo
        image_path_1 = "/var/lib/odoo/modulo/empresa_abogados/static/src/img/imagen_arca_auditores.png"
        run = paragraph.add_run()
        run.add_picture(image_path_1, width=Inches(1))
        paragraph.add_run('  ' * 50)
        image_path_2 = "/var/lib/odoo/modulo/empresa_abogados/static/src/img/arca_auditores_otro.png"
        run = paragraph.add_run()
        run.add_picture(image_path_2, width=Inches(1))
        paragraph.add_run('__'*53).font.color.rgb = RGBColor(255, 140, 0)

        doc.add_paragraph('FACTURA '+str(self.num_factura)).alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph('')

        direccion = self.cliente.direccion if self.cliente.direccion else ""
        doc.add_paragraph(str(self.cliente.name))
        doc.add_paragraph(str(direccion))
        cif = self.cliente.cif if self.cliente.cif else ""
        doc.add_paragraph(str(f'C.I.F.'+ cif))
        doc.add_paragraph('')
        # Add the project name
        doc.add_paragraph('Factura correspondiente a los servicios'+str(self.proyecto.name))

        # Add the invoice details
        doc.add_paragraph('')
        doc.add_paragraph('Base Imponible ..................................... ' + str(self.BI)+'€').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph('% IVA .............................................. ' + str( round(self.importe_IVA, 2)) +'€').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph('TOTAL .............................................. ' + str( round(self.importe_total, 2)) +'€').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph('')
        # Add the signature and date
        doc.add_paragraph('')
        doc.add_paragraph('En Madrid, a fecha y hora de la firma electrónica').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph('')
        doc.add_paragraph('')
        doc.add_paragraph('Fdo. Esther Camacho González').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph('')
        # Add the payment information
        doc.add_paragraph('')
        doc.add_paragraph('Dicha factura la pueden hacer efectiva mediante transferencia a nuestra cuenta corriente del Banco Sabadell (Calle Arturo Soria 316 de Madrid). C.C.C. nº ES4100810638590002108117 a nombre de ARCA AUDITORES S.L.').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Add the data protection information
        doc.add_paragraph('')
        paragraph1 = doc.add_paragraph('De conformidad con lo establecido en la normativa vigente en Protección de Datos de Carácter Personal, le informamos que sus datos serán incorporados al sistema de tratamiento titularidad de ARCA CONSORTIUM SA  con CIF A28911238 y domicilio social sito en CL SANTA ENGRACIA 6 28010, MADRID, con la finalidad de poder remitirle la correspondiente factura. En cumplimiento con la normativa vigente, ARCA CONSORTIUM SA  informa que los datos serán conservados durante EN EL PLAZO LEGALMENTE ESTABLECIDO.')
        paragraph1.runs[0].font.size = Pt(6) 

        # Agregar el segundo párrafo con tamaño de fuente reducido
        paragraph2 = doc.add_paragraph('Podrá ejercer los derechos de acceso, rectificación, limitación de tratamiento, supresión, portabilidad y oposición/revocación, en los términos que establece la normativa vigente en materia de protección de datos, dirigiendo su petición a la dirección postal CL SANTA ENGRACIA 6 28010, MADRID o bien a través de correo electrónico MADRID@ARCACONSORTIUM.NET.')
        paragraph2.runs[0].font.size = Pt(6) 
        
        section = doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]

        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_text = 'ARCA AUDITORES S.L.\n'
        footer_text += 'SEDE CENTRAL SANTA ENGRACIA, 6, 28010 MADRID +34 913192658 madrid@arca-auditores.net\n'
        footer_text +='ARCA AUDITORES, S.L. Inscrito en R.M.Madrid, tomo 38902, folio 90, sección 8, hoja M-691378. CIF: B-88325113'

        # Agregar el texto al párrafo y establecer el tamaño de fuente y color
        run = paragraph.add_run(footer_text)
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(25, 25, 112)  # Color hexadecimal #191970
        run.font.bold=True
        # Save the document
        doc.save('factura.docx')
        # Convert the saved document to binary data
        with open('factura.docx', 'rb') as f:
            binary_data = f.read()

        # Create the attachment record
        attachment = self.env['ir.attachment'].sudo().create({
            'datas': base64.encodestring(binary_data),
            'type': 'binary',
            'res_model': 'empresa_abogados.facturas',
            'res_id': self.proyecto.id,
            'db_datas': 'facturas_Arca_Auditores.docx',
            'name': 'facturas_Arca_Auditores.docx',
        })
        return {
                'type': 'ir.actions.act_url',
                'url': '/web/content/%s?download=true' % attachment.id,
                'target': 'new',
            }