from odoo import models, fields, api
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
import base64
from odoo import models, fields, api
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

class base (models.Model):
    _name = 'empresa_abogados.base'
    _description = 'Modelo base '
    fecha_alta = fields.Date(string='Fecha Alta',default=lambda self: fields.Date.today())
    observaciones = fields.Text(string='Observaciones')
class Proyecto(models.Model):
    _name = 'empresa_abogados.proyecto'
    _description = 'Modelos para almacenar proyectos'
    _inherit = 'empresa_abogados.base'

    id = fields.Char(string="Id", required=True, size=9, default='1')
    tramitacion = fields.Selection([('1', 'Oferta'), ('2', 'Licitacion')], string='Tramitacion')
    empresa = fields.Selection([('1', 'Arca Auditores'), ('2', 'Arca Consortium')], string='Empresa')
    cliente = fields.Many2one('empresa_abogados.clientes', string='Cliente',required=True)
    name = fields.Char(string='Nombre del proyecto',required=True)
    estado = fields.Selection([
    ('en_estudio', 'En Estudio'),
    ('aceptada', 'Aceptada'),
    ('presentado', 'Presentado'),
    ('descartado', 'Descartado'),
    ('adjudicado', 'Adjudicado'),
    ('facturado', 'Facturado'),
    ], string='Estado', compute='_compute_estado_proyecto_pre',store=True, default='en_estudio')

    @api.depends('tramitacion', 'oferta_id.aceptada', 'estudio_id.presentado', 'estudio_id.descartado', 'seguimiento_id.adjudicado', 'gestion_id.facturado')
    def _compute_estado_proyecto_pre(self):
        for proyecto in self:
            if proyecto.tramitacion == '1':
                presentado = any(oferta.aceptada for oferta in proyecto.oferta_id)
                proyecto.estado = 'aceptada' if presentado else 'en_estudio'
            else:
                presentado_estudio = any(estudio.presentado for estudio in proyecto.estudio_id)
                presentado_descartado = any(estudio.descartado for estudio in proyecto.estudio_id)
                presentado_adjudicado = any(seguimiento.adjudicado for seguimiento in proyecto.seguimiento_id)
                presentado_gestion = any(gestion.facturado for gestion in proyecto.gestion_id)

                if presentado_estudio:
                    proyecto.estado = 'presentado'
                elif presentado_descartado:
                    proyecto.estado = 'descartado'
                elif presentado_adjudicado:
                    proyecto.estado = 'adjudicado'
                elif presentado_gestion:
                    proyecto.estado = 'facturado'
                else:
                    proyecto.estado = 'en_estudio'
            

    oferta_id = fields.One2many('empresa_abogados.oferta', 'proyecto_id', string='Oferta')
    estudio_id = fields.One2many('empresa_abogados.estudio', 'proyecto', string='Estudio')
    seguimiento_id = fields.One2many('empresa_abogados.seguimiento', 'proyecto', string='Seguimiento')
    adjudicacion_id = fields.One2many('empresa_abogados.adjudicacion', 'proyecto', string='Adjudicacion')
    facturas_id = fields.One2many('empresa_abogados.facturas', 'proyecto', string='Facturas')
    num_oferta_facturas= fields.Char(related='facturas_id.num_factura',string='Numero de factura')
    gestion_id = fields.One2many('empresa_abogados.gestion','proyectos', string='Gestion')

    BI = fields.Float(string='Base Imponible')
    IVA = fields.Float(string='IVA',default=0.21)
    importe_IVA= fields.Float(compute='calculo_importe_IVA',string='Importe IVA')
    total = fields.Float(compute='calculo_importe',string='Total')
    
  
    documento_ids = fields.One2many('empresa_abogados.documento','proyecto_id',string='Documentos Adjuntos ',ondelete='cascade')   
    equipo = fields.Many2many('res.users', string='Equipo')
    crear_user = fields.Many2one('res.users', 'Usuario creador', readonly=True)
    write_user = fields.Many2one('res.users', 'Último usuario editor', readonly=True)
    
    res_id = fields.Reference([('empresa_abogados.proyecto', 'proyecto')])
    
    attachment_ids = fields.One2many(
        'ir.attachment',
        'res_id',
        string='Adjuntos',
        domain=[('res_model', '=', 'empresa_abogados.proyecto')],
        context={'res_model': 'empresa_abogados.proyecto'},
        ondelete='cascade'
    )
  
    @api.model
    def create(self, vals):
        vals['crear_user'] = self.env.user.id
        return super(Proyecto, self).create(vals)

    def write(self, vals):
        vals['write_user'] = self.env.user.id
        return super(Proyecto, self).write(vals)
    

    @api.depends('BI','IVA','total')
    def calculo_importe(self):
        for elemento in self:
            bi= elemento.BI
            iva= elemento.IVA
            elemento.total= (bi * iva) + bi

    @api.depends('IVA','BI','importe_IVA')
    def calculo_importe_IVA(self):
        for elemento in self:
            bi= elemento.BI
            iva= elemento.IVA
            elemento.importe_IVA= (bi * iva)
   
    duracion= fields.Date(compute='_compute_fecha_duracion',string='fecha duracion',store=True)
   
    @api.depends('oferta_id.fecha_fin_plazo_ejecucion','adjudicacion_id.fecha_fin_plazo_ejecucion','estudio_id.Fin_plazo_presentacion_oferta')
    def _compute_fecha_duracion(self):
        for proyecto in self:
            fecha_pre= proyecto.estudio_id.Fin_plazo_presentacion_oferta
            fecha_f = proyecto.adjudicacion_id.fecha_fin_plazo_ejecucion
            fecha_o = proyecto.oferta_id.fecha_fin_plazo_ejecucion
            if self.tramitacion=='1':
                if fecha_o:
                    proyecto.duracion= fecha_o
                else:
                    proyecto.duracion= fields.Date.today()+ timedelta(days=30)
            else: 
                if fecha_pre:
                    proyecto.duracion = fecha_pre
                    if fecha_f:
                        proyecto.duracion= fecha_f      
                else:
                    proyecto.duracion= fields.Date.today()+ timedelta(days=30)
    fecha_alarma = fields.Date(compute='_compute_fecha_alarma', store=True)
    @api.depends('oferta_id.fecha_fin_plazo_ejecucion','adjudicacion_id.fecha_fin_plazo_ejecucion','estudio_id.fecha_presentacion', 'duracion')
    def _compute_fecha_alarma(self):
        for proyecto in self:
            fecha_pre = proyecto.estudio_id.Fin_plazo_presentacion_oferta
            fecha_f = proyecto.adjudicacion_id.fecha_fin_plazo_ejecucion
            fecha_o = proyecto.oferta_id.fecha_fin_plazo_ejecucion
            if self.tramitacion =='1':
                if fecha_o:
                    proyecto.fecha_alarma = fecha_o - timedelta(days=3)
                else:
                    proyecto.fecha_alarma = fields.Date.today()+ timedelta(days=30)
            else:
                if fecha_pre:
                    proyecto.fecha_alarma  = fecha_pre - timedelta(days=3)
                    if fecha_f:
                        proyecto.fecha_alarma  = fecha_f - timedelta(days=5)
                else:
                    proyecto.fecha_alarma = False

            #if proyecto.fecha_alarma and proyecto.fecha_alarma == fields.Date.today():
            #   proyecto.create_notification()
 

    def create_notification(self):
        message1='FECHAS LIMITE CERCA'
        message='Proyectos: '
      
        proyectos = self.env['empresa_abogados.proyecto'].search([('fecha_alarma', '=', fields.Date.today())])
        proyecto_names = [proyecto.name for proyecto in proyectos]
        proyecto_names_str = ", ".join(proyecto_names)

        if proyecto_names:
            return {
                'type': 'ir.actions.client',
                'tag': 'display_notification',
                'params': {
                    'title': message1,
                    'message': f'{message} {proyecto_names_str}',
                    'type': 'warning',
                    'sticky': True,
                },
            }
        else:
            return False
    
    #campos para el informe 
    fecha_duracion = fields.Char(compute='_compute_fecha_anios')
    importe_facturas= fields.Char(compute='_compute_importe_facturas_suma')
    @api.depends('gestion_id.fecha_finalizacion', 'fecha_alta')
    def _compute_fecha_anios(self):
        for proyecto in self:

            fecha = proyecto.fecha_alta
            gestion_fecha= None
            for gestion in proyecto.gestion_id:
                if gestion.fecha_finalizacion:
                    gestion_fecha = gestion.fecha_finalizacion
            if gestion_fecha :
                proyecto.fecha_duracion = str( gestion_fecha.year)
                if fecha:
                     proyecto.fecha_duracion = None
                     proyecto.fecha_duracion = str( gestion_fecha.year)+" - " + str(fecha.year)
                     if gestion_fecha.year == fecha.year:
                         proyecto.fecha_duracion = None
                         proyecto.fecha_duracion = str( gestion_fecha.year)
            else:
                proyecto.fecha_duracion = ''

    @api.depends('facturas_id.importe_total')
    def _compute_importe_facturas_suma(self):
        for proyecto in self:
            importe_total = sum(factura.importe_total for factura in proyecto.facturas_id)
            proyecto.importe_facturas = str(importe_total) if importe_total else ''
   
    def generar_certificado(self):
            doc = Document()
            #header   
            paragraph = doc.add_paragraph("MEMBRETE DEL ORGANISMO")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.style = "Heading2"

            # Obtener el objeto Run
            run = paragraph.runs[0]

            # Establecer el tamaño del texto
            run.font.size = Pt(14)  # Ajusta el valor según tus necesidades
            run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
            doc.add_paragraph('')
            #content
            content_paragraph = doc.add_paragraph()
            content_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            content_paragraph.add_run("D.(NOMBRE DE LA PERSONA QUE FIRMA), como (CARGO QUE OCUPA DICHA PERSONA) de ")

            cliente = self.cliente.name if self.cliente else ""
            content_paragraph.add_run(self.cliente.name).bold = True

            content_text = ", entidad contratante del servicio "
            content_paragraph.add_run(content_text)

            project_name = self.name if self.name else ""
            content_paragraph.add_run(self.name).bold = True

            # Add the "CERTIFICA" section
            certifica_paragraph = doc.add_paragraph('')
            certifica_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            certifica_text = "CERTIFICA:"
            certifica_run = certifica_paragraph.add_run(certifica_text)
            certifica_run.bold = True
            # Add the details section
            details_paragraph = doc.add_paragraph()
            details_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            details_paragraph.add_run("Que la empresa ")

            empresa = self.empresa 
            if self.empresa :
                if empresa=='1':
                    empresa='Arca Auditores'
                else: 
                    empresa='Arca Consortium'
            else:
                empresa = ""
            details_paragraph.add_run(empresa).bold = True

            details_text = ", adjudicataria del referido servicio, ejecutó los trabajos de acuerdo a las condiciones y a satisfacción del "
            details_paragraph.add_run(details_text)

            details_paragraph.add_run(cliente).bold = True

            details_text = " durante el año "
            details_paragraph.add_run(details_text)

            fecha_duracion = self.fecha_duracion if self.fecha_duracion else ""
            details_paragraph.add_run(fecha_duracion).bold = True

            details_paragraph.add_run(", siendo el importe del contrato de ")

            importe_facturas = self.importe_facturas if self.importe_facturas else ""
            details_paragraph.add_run(importe_facturas).bold = True

            details_paragraph.add_run(" euros (IVA incluido), y las personas que ejecutaron dicho servicio fueron:")

            # Add the list of persons
            persons_paragraph = doc.add_paragraph()
            persons_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            persons_list = [
                "Esther Camacho",
                "Jaime García-Rosado",
                "Luis Felipe Suárez-Olea"
            ]

            for person in persons_list:
                persons_paragraph.add_run(f"• {person}\n")

            # Add the date
            date_paragraph = doc.add_paragraph()
            date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            date_text = "Y para que así conste, y a quien pudiera interesar, expido el presente certificado en Madrid, a "
            date_paragraph.add_run(date_text)

            date = fields.datetime.now().strftime('%d/%m/%Y')
            date_paragraph.add_run(date).bold = True

            # Add the signature and footer
            footer_paragraph = doc.add_paragraph()
            footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            footer_text = "CARGO DE LA PERSONA QUE FIRMA\n\n\n\n"
            footer_paragraph.add_run(footer_text)
            
            doc.add_paragraph('')
            footer_paragraph.add_run("FIRMA Y SELLO\n\n\n\n")
            
            doc.add_paragraph('')
            footer_paragraph.add_run("NOMBRE DE LA PERSONA QUE FIRMA")
            
            doc.add_paragraph('')
            # Save the document
            doc.save('certificado.docx')
            # Save the document
       
            # Convert the saved document to binary data
            with open('certificado.docx', 'rb') as f:
                binary_data = f.read()

            # Create the attachment record
            attachment = self.env['ir.attachment'].sudo().create({
                'datas': base64.encodestring(binary_data),
                'type': 'binary',
                'res_model': 'empresa_abogados.proyecto',
                'res_id': self.id,
                'db_datas': str(self.name)+'.docx',
                'name': str(self.name)+'.docx',
            })
            return {
                'type': 'ir.actions.act_url',
                'url': '/web/content/%s?download=true' % attachment.id,
                'target': 'new',
            }